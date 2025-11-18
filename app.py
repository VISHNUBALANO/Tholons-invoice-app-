import streamlit as st
import pandas as pd
import json
import os
import re
import math
from datetime import date, datetime, timedelta
from docxtpl import DocxTemplate
from io import BytesIO
from num2words import num2words
from dateutil.parser import isoparse
from typing import Dict, Any, List, Optional, Tuple, Callable
from pymongo import MongoClient

# --- Configuration ---
DATA_DIR = '_data'
TEMPLATES_DIR = 'templates'
INVOICES_DIR = 'generated_invoices'
EMPLOYEES_FILE = os.path.join(DATA_DIR, 'employees.json')  # no longer used, kept for compatibility
INVOICES_FILE = os.path.join(DATA_DIR, 'invoices.json')   # no longer used
META_FILE = os.path.join(DATA_DIR, 'meta.json')           # no longer used
TEMPLATE_PATH = os.path.join(TEMPLATES_DIR, 'invoice.docx')

DEFAULT_META = {"invoice_prefix": "TH-MD-", "last_number": 0}

# --- MongoDB helpers ---

@st.cache_resource(show_spinner=False)
def get_db():
    """Return a cached MongoDB database connection."""
    client = MongoClient(st.secrets["MONGO_URI"])
    return client[st.secrets["MONGO_DB"]]

def mongo_fetch_all(collection_name: str) -> List[dict]:
    """Fetch all documents from a collection and strip Mongo _id."""
    db = get_db()
    col = db[collection_name]
    docs: List[dict] = []
    for doc in col.find({}):
        doc.pop("_id", None)
        docs.append(doc)
    return docs

# --- File and misc utilities (JSON loaders kept but no longer used for persistence) ---

def ensure_dirs():
    """Ensure required directories exist."""
    for d in [DATA_DIR, TEMPLATES_DIR, INVOICES_DIR]:
        os.makedirs(d, exist_ok=True)

def load_json(filepath, default_data=None):
    """Legacy JSON loader, no longer used with Mongo, kept for safety."""
    try:
        if not os.path.exists(filepath):
            raise FileNotFoundError
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return data if data is not None else default_data
    except FileNotFoundError:
        save_json(filepath, default_data)
        return default_data
    except json.JSONDecodeError:
        st.error(f"Error decoding JSON from {os.path.basename(filepath)}. File might be corrupt. Returning default.")
        save_json(filepath, default_data)
        return default_data
    except Exception as e:
        st.error(f"An unexpected error occurred loading {os.path.basename(filepath)}: {e}")
        save_json(filepath, default_data)
        return default_data

def save_json(filepath, data):
    """Legacy JSON saver, no longer used with Mongo, kept for safety."""
    try:
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2)
    except Exception as e:
        st.error(f"Error saving data to {os.path.basename(filepath)}: {e}")

def get_next_id(data):
    """Get the next sequential ID for a list of records."""
    if not data:
        return 1
    return max(item.get('id', 0) for item in data) + 1

def safe_filename(name):
    """Sanitize a string for use as a Windows filename."""
    s = re.sub(r'[\\/:*?"<>|\r\n]', '_', str(name))
    s = re.sub(r'_{2,}', '_', s)
    s = s.strip().strip('.')
    return s

def get_ordinal_date(date_obj: date, include_year: bool = True) -> str:
    """Converts a date object to the 'Dth Month YYYY' format (for example, 19th September 2025)."""
    if date_obj is None:
        return ""

    def suffix(day):
        if 11 <= day <= 13:
            return 'th'
        return {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')

    day = date_obj.day
    if isinstance(date_obj, datetime):
        date_obj = date_obj.date()

    date_format = f"%d{suffix(day)} %B"
    if include_year:
        date_format += " %Y"

    return date_obj.strftime(date_format).replace(' 0', ' ')

def format_invoice_no(prefix: str, number_suffix: str) -> str:
    """Generate the full invoice number string in TH-MD-YYYY-XX format."""
    current_year = datetime.now().year
    return f"{prefix}{current_year}-{number_suffix}"

def amount_to_words(amount):
    """Convert a numeric amount (USD) to English words, removing ', Zero Cents'."""
    try:
        amount = round(amount, 2)
        words = num2words(amount, to='currency', currency='USD').title()

        words = re.sub(r', Zero Cents', '', words)
        words = re.sub(r' And Zero Cents', '', words)
        words = re.sub(r' And [A-Za-z\s]+ Cents', '', words)

        if 'Dollars' in words:
            words = words.replace('Dollars', 'Dollars Only')
        elif 'Dollar' in words:
            words = words.replace('Dollar', 'Dollar Only')

        return words

    except Exception:
        return f"Amount in words conversion failed ({amount:.2f})"

def create_starter_template(path):
    """Create a minimal DOCX template with required placeholders if one does not exist."""
    if os.path.exists(path):
        return

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        doc.add_heading("Addison Weekly Invoice (STARTER TEMPLATE - FIX ME)", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Invoice No: {{{{ invoice_no }}}} | Dated: {{{{ invoice_date }}}}")
        doc.add_paragraph(f"Bill To Name: {{{{ bill_to_name }}}}\nAddress: {{{{ bill_to_address }}}}")

        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Employee/Role'
        hdr_cells[1].text = 'Rate'
        hdr_cells[2].text = 'Hours'
        hdr_cells[3].text = 'Amount (salary)'

        data_cells = table.rows[1].cells
        data_cells[0].text = f"{{{{ name }}}} ({{{{ role }}}}) at {{{{ client }}}}"
        data_cells[1].text = f"{{{{ rates }}}}"
        data_cells[2].text = f"{{{{ hours }}}}"
        data_cells[3].text = f"{{{{ salary }}}}"

        doc.add_paragraph("--- TOTAL ---")
        doc.add_paragraph(f"Total Amount Due for Week of {{{{ before_invoice_date }}}} : USD {{{{ salary }}}} (In words: {{{{ salary_words }}}})")

        doc.save(path)
        st.warning(f"Template was missing. A starter template was created at `{path}`. Replace this file with your custom template.")
    except ImportError:
        st.error("Cannot create starter template. Please ensure 'python-docx' is installed.")
    except Exception as e:
        st.error(f"Error creating starter template: {e}")

# --- Data load and save using MongoDB ---

@st.cache_data(show_spinner=False)
def load_data():
    """
    Load all persistent data from MongoDB. Caches this data until explicitly cleared.
    """
    ensure_dirs()

    if not os.path.exists(TEMPLATE_PATH):
        create_starter_template(TEMPLATE_PATH)

    db = get_db()
    employees_col = db["employees"]
    invoices_col = db["invoices"]
    meta_col = db["meta"]

    # Seed default employee if collection is empty
    default_employees = [{
        "id": 1,
        "name": "Cory Caine",
        "role": "Salesforce Architect",
        "rate": 104.00,
        "client": "TPx Communications",
        "bill_to_name": "Accounts Payable",
        "bill_to_address": "Attn: Lana Buchbinder, Director - Recruiting\nMondo, 102 Madison, 7th Floor\nNew York, NY 10016",
        "created_at": date.today().isoformat()
    }]

    if employees_col.count_documents({}) == 0:
        employees_col.insert_many(default_employees)

    # Fetch employees
    employees = mongo_fetch_all("employees")

    # Fetch invoices
    invoices = mongo_fetch_all("invoices")

    # Fetch meta
    meta_doc = meta_col.find_one({"_id": "meta"})
    if meta_doc is None:
        meta = DEFAULT_META.copy()
        meta_col.insert_one({"_id": "meta", **meta})
    else:
        meta = {
            "invoice_prefix": meta_doc.get("invoice_prefix", DEFAULT_META["invoice_prefix"]),
            "last_number": meta_doc.get("last_number", DEFAULT_META["last_number"])
        }

    employee_dict = {emp['name']: emp for emp in employees}

    return employees, employee_dict, invoices, meta

def save_employee_data(employees):
    """
    Save the updated employees list to MongoDB and clear the cache to reload the data.
    """
    db = get_db()
    employees_col = db["employees"]
    employees_col.delete_many({})
    if employees:
        employees_col.insert_many(employees)
    load_data.clear()
    st.rerun()

def save_invoice_data(invoices, new_meta):
    """Save the updated invoices list and meta data to MongoDB."""
    db = get_db()
    invoices_col = db["invoices"]
    meta_col = db["meta"]

    invoices_col.delete_many({})
    if invoices:
        invoices_col.insert_many(invoices)

    meta_col.update_one({"_id": "meta"}, {"$set": new_meta}, upsert=True)

    st.session_state.invoices_data = invoices
    st.session_state.meta_data = new_meta

# --- Invoice rendering ---

def render_and_save_invoice(context, invoice_no, employee_name):
    """Render the invoice using docxtpl and save it to a temporary buffer."""
    try:
        tpl = DocxTemplate(TEMPLATE_PATH)
        tpl.render(context)

        file_stream = BytesIO()
        tpl.save(file_stream)
        file_stream.seek(0)

        safe_emp_name = safe_filename(employee_name)
        safe_inv_no = safe_filename(invoice_no)
        filename = f"{safe_inv_no}_{safe_emp_name}.docx"
        generated_path = os.path.join(INVOICES_DIR, filename)

        with open(generated_path, 'wb') as f:
            f.write(file_stream.getbuffer())

        return file_stream, generated_path

    except Exception as e:
        if "jinja2" in str(e).lower() or "variable" in str(e).lower() or "template" in str(e).lower():
            st.error(f"Template Rendering Error: Check your template placeholders and context data. Details: {e}")
        elif "No such file or directory" in str(e):
            st.error(f"Template file not found at: {TEMPLATE_PATH}. Please check the file path.")
        else:
            st.error(f"An unexpected error occurred during rendering: {e}")
        return None, None

def convert_docx_to_pdf_stream(docx_stream: BytesIO) -> Optional[BytesIO]:
    """
    Placeholder function for DOCX to PDF conversion.
    To enable, wire up docx2pdf or LibreOffice tooling and return a BytesIO.
    """
    st.warning("PDF conversion is unavailable in this environment. Download the DOCX and convert locally.")
    return None

# --- Streamlit UI ---

def main():
    """Main Streamlit application function."""
    st.set_page_config(layout="wide", page_title="Addison Invoice Generator")
    st.title("💰 Addison Weekly Invoice Generator")

    # Load Initial Data
    if 'employees_data' not in st.session_state:
        employees, employee_dict, invoices, meta = load_data()
        st.session_state.employees_data = employees
        st.session_state.employee_dict = employee_dict
        st.session_state.invoices_data = invoices
        st.session_state.meta_data = meta
        st.session_state.download_data_docx = None
        st.session_state.download_data_pdf = None
        st.session_state.download_filename_base = None
        st.session_state.download_success = False
        st.session_state.last_invoice_no = None

    employees = st.session_state.employees_data
    employee_dict = st.session_state.employee_dict
    invoices = st.session_state.invoices_data
    meta = st.session_state.meta_data

    with st.sidebar:
        st.header("👤 Employee Registry")

        with st.form("employee_form"):
            st.subheader("Add/Edit Employee")

            employee_names = ["-- New Employee --"] + list(employee_dict.keys())
            selected_name = st.selectbox("Select Employee to Edit", options=employee_names, key="edit_employee_select")

            initial_data = {}
            if selected_name != "-- New Employee --":
                initial_data = employee_dict.get(selected_name, {})

            default_name = initial_data.get('name', '')
            default_role = initial_data.get('role', '')
            default_client = initial_data.get('client', '')
            default_rate = initial_data.get('rate', 100.00)
            default_bill_to_name = initial_data.get('bill_to_name', '')
            default_bill_to_address = initial_data.get('bill_to_address', '')

            name = st.text_input("Name", value=default_name, key="emp_name")
            role = st.text_input("Role", value=default_role, key="emp_role")
            client = st.text_input("Client", value=default_client, key="emp_client")
            rate = st.number_input("Default Rate (USD/hr)", min_value=0.0, step=0.01, format="%.2f",
                                   value=default_rate, key="emp_rate")
            bill_to_name = st.text_input("Bill To Name", value=default_bill_to_name, key="emp_bill_to_name")
            bill_to_address = st.text_area("Bill To Address", value=default_bill_to_address, key="emp_bill_to_address")

            submitted = st.form_submit_button("Save Employee")

            if submitted:
                if not name:
                    st.error("Employee Name is required.")
                    return
                elif selected_name != "-- New Employee --" and name != selected_name and name in employee_dict:
                    st.error(f"An employee with the name '{name}' already exists.")
                    return
                else:
                    new_employee = {
                        "name": name,
                        "role": role,
                        "rate": rate,
                        "client": client,
                        "bill_to_name": bill_to_name,
                        "bill_to_address": bill_to_address,
                        "created_at": date.today().isoformat()
                    }

                    if selected_name == "-- New Employee --":
                        new_employee["id"] = get_next_id(employees)
                        employees.append(new_employee)
                        st.success(f"Employee '{name}' added successfully!")
                    else:
                        index_to_update = next((i for i, emp in enumerate(employees) if emp['name'] == selected_name), -1)
                        if index_to_update != -1:
                            new_employee['id'] = employees[index_to_update]['id']
                            employees[index_to_update] = new_employee
                            st.success(f"Employee '{name}' updated successfully!")
                        else:
                            st.error("Error: Could not find employee to update.")

                    save_employee_data(employees)

        st.subheader("Current Employees")
        if employees:
            df_employees = pd.DataFrame(employees)
            df_display = df_employees[['name', 'role', 'client', 'rate']].rename(columns={'rate': 'Rate (USD/hr)'})
            st.dataframe(df_display, use_container_width=True, hide_index=True)
        else:
            st.info("No employees registered yet.")

    # Main Area: Tabs
    tab1, tab2 = st.tabs(["🧾 Generate Invoice", "📊 Invoice Report"])

    # Tab 1: Invoice Creation
    with tab1:
        st.header("Invoice Creation")

        employee_options = list(employee_dict.keys())
        if not employee_options:
            st.error("Please add at least one employee in the sidebar to generate an invoice.")
            st.session_state.download_success = False
            return

        selected_employee_name = st.selectbox("Select Employee for Invoice", options=employee_options, key="inv_employee")
        employee_details = employee_dict.get(selected_employee_name)

        if employee_details:
            st.caption("Employee and Client Details")
            col1, col2, col3 = st.columns(3)
            col1.info(f"**Role:** {employee_details['role']}")
            col2.info(f"**Client:** {employee_details['client']}")
            col3.info(f"**Default Rate:** ${employee_details['rate']:.2f}/hr")

            st.caption("Bill To Details")
            st.code(f"**{employee_details['bill_to_name']}**\n{employee_details['bill_to_address']}")

        st.markdown("---")

        with st.form("invoice_form"):
            st.subheader("Invoice Details")

            col_inv_date, col_week_of = st.columns(2)

            # 1. Invoice Date
            invoice_date_selected = col_inv_date.date_input(
                "📅 Select Invoice Date",
                value=datetime.now().date(),
                key="form_invoice_date"
            )

            # 2. Week Of Date
            week_of_date_default = invoice_date_selected - timedelta(days=5)
            week_of_date_selected = col_week_of.date_input(
                "🗓️ Select Week Of Date",
                value=week_of_date_default,
                key="form_week_of_date"
            )

            # 3. Formatted tags
            invoice_date_formatted = get_ordinal_date(invoice_date_selected, include_year=True)
            week_of_date_formatted = get_ordinal_date(week_of_date_selected, include_year=False)

            st.markdown(f"**Formatted Invoice Date ({{{{ invoice_date }}}}):** `{invoice_date_formatted}`")
            st.markdown(f"**Formatted Week Of Date ({{{{ before_invoice_date }}}}):** `{week_of_date_formatted}`")
            st.markdown("---")

            # Manual Dropdown for Invoice Number Suffix (01..200)
            number_options = [str(i).zfill(2) for i in range(1, 201)]
            selected_invoice_suffix = st.selectbox(
                f"Select Invoice Number Suffix (TH-MD-{datetime.now().year}-XX)",
                options=number_options,
                key="form_invoice_suffix",
                index=0
            )

            # Construct invoice number
            new_inv_no = format_invoice_no(meta['invoice_prefix'], selected_invoice_suffix)
            st.markdown(f"**Full Invoice No. ({{{{ invoice_no }}}}):** `{new_inv_no}`")

            col_hours, col_rate, col_client = st.columns(3)

            hours = col_hours.number_input("Working Hours", min_value=0.0, step=0.5, format="%.2f", value=40.0, key="form_hours")

            rate_override = col_rate.number_input("Rate Override (USD/hr)", min_value=0.0, step=0.01, format="%.2f",
                                                  value=employee_details['rate'],
                                                  key=f"form_rate_{employee_details['id']}")

            client_override = col_client.text_input("Client Override", value=employee_details['client'],
                                                    key=f"form_client_{employee_details['id']}")

            # Salary calc
            amount = round(hours * rate_override, 2)
            st.metric(label="Total Amount Due", value=f"${amount:,.2f}")
            amount_words_str = amount_to_words(amount)
            st.caption(f"Amount in Words: *{amount_words_str}*")

            generate_button = st.form_submit_button("Generate & Save Invoice")

            if generate_button:
                st.session_state.download_success = False

                context = {
                    'invoice_no': new_inv_no,
                    'invoice_date': invoice_date_formatted,
                    'bill_to_name': employee_details['bill_to_name'],
                    'bill_to_address': employee_details['bill_to_address'],
                    'before_invoice_date': week_of_date_formatted,
                    'name': employee_details['name'],
                    'role': employee_details['role'],
                    'client': client_override,
                    'rates': f"{rate_override:.2f}",
                    'hours': f"{hours:.2f}",
                    'salary': f"{amount:,.2f}",
                    'salary_words': amount_words_str,
                }

                file_stream_docx, generated_path = render_and_save_invoice(context, new_inv_no, employee_details['name'])

                if file_stream_docx:
                    new_invoice = {
                        "id": get_next_id(invoices),
                        "employee_id": employee_details['id'],
                        "employee_name": employee_details['name'],
                        "invoice_no": new_inv_no,
                        "invoice_date": invoice_date_selected.isoformat(),
                        "hours": hours,
                        "rate": rate_override,
                        "amount": amount,
                        "client": client_override,
                        "generated_path": generated_path,
                        "created_at": date.today().isoformat()
                    }
                    invoices.append(new_invoice)

                    new_meta = {"invoice_prefix": meta['invoice_prefix'], "last_number": meta['last_number']}
                    save_invoice_data(invoices, new_meta)

                    # Download filename base
                    date_for_filename = invoice_date_selected.strftime("%d-%b").upper()
                    safe_emp_name = safe_filename(employee_details['name']).replace(' ', '_')
                    base_filename = f"Tholons_Addision_Invoice_{safe_emp_name}_{date_for_filename}"

                    # Try PDF conversion
                    file_stream_docx.seek(0)
                    file_stream_pdf = convert_docx_to_pdf_stream(file_stream_docx)

                    # Save into session
                    file_stream_docx.seek(0)
                    st.session_state.download_data_docx = file_stream_docx
                    st.session_state.download_data_pdf = file_stream_pdf
                    st.session_state.download_filename_base = base_filename
                    st.session_state.download_success = True
                    st.session_state.last_invoice_no = new_inv_no

                    st.rerun()

        # Success block outside the form
        if st.session_state.download_success:
            st.success(f"Invoice {st.session_state.last_invoice_no} for {selected_employee_name} generated successfully and saved to `{INVOICES_DIR}`!")

            download_cols = st.columns(2)

            # DOCX bytes explicit
            docx_bytes = (
                st.session_state.download_data_docx.getvalue()
                if hasattr(st.session_state.download_data_docx, "getvalue")
                else st.session_state.download_data_docx
            )

            download_cols[0].download_button(
                label=f"Download {st.session_state.download_filename_base}.docx (Word)",
                data=docx_bytes,
                file_name=f"{st.session_state.download_filename_base}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_invoice_docx_btn",
                type="primary",
            )

            # PDF only if we have bytes
            pdf_bytes = None
            if st.session_state.download_data_pdf is not None:
                pdf_stream = st.session_state.download_data_pdf
                pdf_bytes = pdf_stream.getvalue() if hasattr(pdf_stream, "getvalue") else pdf_stream

            if pdf_bytes:
                download_cols[1].download_button(
                    label=f"Download {st.session_state.download_filename_base}.pdf (PDF)",
                    data=pdf_bytes,
                    file_name=f"{st.session_state.download_filename_base}.pdf",
                    mime="application/pdf",
                    key="download_invoice_pdf_btn",
                    type="secondary",
                )
            else:
                download_cols[1].button(
                    label=f"Download {st.session_state.download_filename_base}.pdf (PDF)",
                    disabled=True,
                    help="PDF conversion is unavailable. Download the DOCX and convert it locally.",
                    key="download_invoice_pdf_unavailable",
                )

    # Tab 2: Invoice report
    with tab2:
        st.header("Generated Invoice Report")

        if invoices:
            df_invoices = pd.DataFrame(invoices)

            df_display = df_invoices.rename(columns={
                'id': 'ID', 'invoice_no': 'Invoice No', 'invoice_date': 'Date',
                'employee_name': 'Employee', 'client': 'Client', 'hours': 'Hours',
                'rate': 'Rate ($)', 'amount': 'Amount ($)', 'generated_path': 'File Path'
            })

            df_display['Rate ($)'] = df_display['Rate ($)'].apply(lambda x: f"${x:,.2f}")
            df_display['Amount ($)'] = df_display['Amount ($)'].apply(lambda x: f"${x:,.2f}")
            df_display['Hours'] = df_display['Hours'].apply(lambda x: f"{x:,.2f}")

            df_display = df_display.sort_values(by='ID', ascending=False)

            st.dataframe(df_display, use_container_width=True, hide_index=True)

            @st.cache_data
            def convert_df_to_csv(df):
                return df.to_csv(index=False).encode('utf-8')

            csv_data = convert_df_to_csv(df_invoices)

            st.download_button(
                label="Export Report to CSV",
                data=csv_data,
                file_name=f"Addison_Invoices_Report_{date.today().isoformat()}.csv",
                mime='text/csv',
                key="export_csv_btn"
            )

        else:
            st.info("No invoices have been generated yet.")

if __name__ == "__main__":
    # Ensure all download state variables are initialized
    if 'download_success' not in st.session_state:
        st.session_state.download_data_docx = None
        st.session_state.download_data_pdf = None
        st.session_state.download_filename_base = None
        st.session_state.download_success = False
        st.session_state.last_invoice_no = None

    main()
